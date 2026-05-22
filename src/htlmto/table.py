import re
from docx import Document
from docx.shared import Pt
from collections import defaultdict


def add_mixed_text(paragraph, parts, clear=True):
    if clear:
        paragraph.clear()
    for part, formatting in parts:

        if "bar" in formatting:
            bar_run = paragraph.add_run(f"{part}")

            bar_run = paragraph.add_run("\u0305")
            bar_run.font.superscript = True
            return
        run = paragraph.add_run(part)

        if "it" in formatting:
            run.italic = True
        if "sub" in formatting:
            run.font.subscript = True
        if "bold" in formatting:
            run.font.bold = True
        if "sup" in formatting:
            run.font.superscript = True


def same_elements(s1, s2, s3, s4):
    set1 = set(
        [
            re.sub(r"\d+$", "", str(s1)).strip(),
            re.sub(r"\d+$", "", str(s2)).strip(),
        ]
    )
    set2 = set(
        [
            re.sub(r"\d+$", "", str(s3)).strip(),
            re.sub(r"\d+$", "", str(s4)).strip(),
        ]
    )

    if len(set1):
        return set1 == set2
    return False


def create_word_table(site_conns, pair_cohp_vals, sample_name="SAMPLE_NAME"):
    condensed_dict = {}

    for site, conns in site_conns.items():
        counts = defaultdict(int)

        for neighbor, distance in conns:
            rounded_dist = round(distance, 4)
            key = (neighbor, rounded_dist)
            counts[key] += 1

        condensed_interactions = []
        for neighbor, dist in counts.keys():
            count = counts[(neighbor, dist)]
            condensed_interactions.append((neighbor, dist, count))

        condensed_interactions.sort(key=lambda x: (x[1], x[0]))

        condensed_dict[site] = condensed_interactions
    # for k, v in condensed_dict.items():
    #     print(k, v)

    doc = Document()
    total_entries = sum([len(v) for v in condensed_dict.values()])

    # Coord table
    doc.add_paragraph(f"\nTable #. Coordination Environments in {sample_name}")
    table = doc.add_table(rows=total_entries + 1, cols=6)
    table.style = "Table Grid"
    style = table.style
    style.font.size = Pt(12)
    table.autofit = True

    header_row = table.rows[0]
    add_mixed_text(header_row.cells[0].paragraphs[0], [["atom", "bold"]])
    header_row.cells[1].text = " "
    add_mixed_text(header_row.cells[2].paragraphs[0], [["count", "bold"]])
    add_mixed_text(
        header_row.cells[3].paragraphs[0], [["d", "it-bold"], [" (Å)", "bold"]]
    )
    add_mixed_text(
        header_row.cells[4].paragraphs[0], [["-COHP (per atom)", "bold"]]
    )
    add_mixed_text(
        header_row.cells[5].paragraphs[0], [["-ICOHP (per atom)", "bold"]]
    )

    for i in range(6):
        header_row.cells[i].bold = True

    target_row_idx = 0
    i = 0
    for site, neighbors in condensed_dict.items():
        cohp_sum = []
        first_row = None
        for j, entry in enumerate(neighbors):
            target_row_idx = i + 1
            col_offset = 0
            row = table.rows[target_row_idx]

            if j == 0:
                first_row = row  # row.cells[0 + col_offset].text = str(site)
            if j == 1:
                row.cells[0 + col_offset].text = "CN " + str(
                    sum([v[-1] for v in neighbors])
                )
            row.cells[1 + col_offset].text = str(entry[0])
            row.cells[2 + col_offset].text = f"{entry[2]}\u00d7"
            row.cells[3 + col_offset].text = str(entry[1])

            val = [
                (v["cohp_val"], v["icohp_val"], abs(v["d"] - entry[1]))
                for v in pair_cohp_vals
                if same_elements(v["s1"], v["s2"], site, entry[0])
                and v["count"] == entry[2]
            ]
            val = sorted(val, key=lambda x: x[1])
            val = val[0]

            cohp_sum.append(round(val[0], 2))

            row.cells[4 + col_offset].text = f"{val[0]/entry[2]:2.2f}"
            row.cells[5 + col_offset].text = f"{val[1]/entry[2]:2.2f}"
            i += 1

        first_row.cells[0 + col_offset].text = f"{site} ({sum(cohp_sum):.2f})"
        # B.O. = sum of icohp, wighted

    doc.save("cohp.docx")


if __name__ == "__main__":
    d = {
        "Er1": [
            ("Co1", 2.708),
            ("Co1", 2.838),
            ("Co1", 2.838),
            ("Co2", 3.118),
            ("Co2", 3.118),
            ("Co2", 3.118),
            ("Co2", 3.118),
            ("In1", 3.244),
            ("In1", 3.244),
            ("In1", 3.244),
            ("In1", 3.244),
            ("In1", 3.255),
            ("In1", 3.255),
        ],
        "In1": [
            ("Co1", 2.71),
            ("Co1", 2.71),
            ("Co2", 2.791),
            ("Co2", 2.791),
            ("Er1", 3.244),
            ("Er1", 3.244),
            ("Er1", 3.244),
            ("Er1", 3.244),
            ("Er1", 3.255),
            ("Er1", 3.255),
            ("In1", 3.256),
            ("In1", 3.256),
        ],
        "Co1": [
            ("Co2", 2.467),
            ("Co2", 2.467),
            ("Co2", 2.467),
            ("Co2", 2.467),
            ("Er1", 2.708),
            ("In1", 2.71),
            ("In1", 2.71),
            ("Er1", 2.838),
            ("Er1", 2.838),
            ("Co1", 2.848),
            ("Co1", 2.848),
        ],
        "Co2": [
            ("Co1", 2.467),
            ("Co1", 2.467),
            ("Co1", 2.467),
            ("Co1", 2.467),
            ("Co2", 2.5),
            ("Co2", 2.5),
            ("In1", 2.791),
            ("In1", 2.791),
            ("Er1", 3.118),
            ("Er1", 3.118),
            ("Er1", 3.118),
            ("Er1", 3.118),
        ],
    }

    cohp = [
        {
            "s1": "In1",
            "s2": "In1",
            "count": 2,
            "d": 3.256,
            "cohp_val": 0.05695581345177665,
            "icohp_val": 2.0113838657994925,
        },
        {
            "s1": "In1",
            "s2": "Er1",
            "count": 4,
            "d": 3.244,
            "cohp_val": 0.33924195145939084,
            "icohp_val": 2.435866217322335,
        },
        {
            "s1": "In1",
            "s2": "Er1",
            "count": 2,
            "d": 3.255,
            "cohp_val": 0.20274597271573605,
            "icohp_val": 1.4759492753807106,
        },
        {
            "s1": "In1",
            "s2": "Co1",
            "count": 2,
            "d": 2.71,
            "cohp_val": 0.020516779822335025,
            "icohp_val": 2.503895467639594,
        },
        {
            "s1": "In1",
            "s2": "Co2",
            "count": 2,
            "d": 2.791,
            "cohp_val": -0.056406457170050765,
            "icohp_val": 2.02955345177665,
        },
        {
            "s1": "Er1",
            "s2": "Co1",
            "count": 1,
            "d": 2.708,
            "cohp_val": 0.07459132043147208,
            "icohp_val": 0.7612863807106599,
        },
        {
            "s1": "Er1",
            "s2": "Co1",
            "count": 2,
            "d": 2.838,
            "cohp_val": 0.1571200948604061,
            "icohp_val": 1.460398782677665,
        },
        {
            "s1": "Er1",
            "s2": "Co2",
            "count": 4,
            "d": 3.118,
            "cohp_val": 0.22468874587563453,
            "icohp_val": 1.974004094860406,
        },
        {
            "s1": "Co2",
            "s2": "Co2",
            "count": 4,
            "d": 2.467,
            "cohp_val": -0.19260991275380712,
            "icohp_val": 4.938165917195431,
        },
        {
            "s1": "Co2",
            "s2": "Co2",
            "count": 2,
            "d": 2.499,
            "cohp_val": -0.12262505615482233,
            "icohp_val": 2.4836235942258886,
        },
        {
            "s1": "Co2",
            "s2": "Co1",
            "count": 2,
            "d": 2.848,
            "cohp_val": -0.0420850453680203,
            "icohp_val": 1.388913279505076,
        },
    ]

    create_word_table(d, cohp)
